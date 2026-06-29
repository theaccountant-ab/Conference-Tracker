"""Process call-for-papers files submitted by conference organizers.

This is the "submission" input channel for the hybrid model: organizers submit
their CFP (PDF / Word / text) via the site's "Submit your CFP" button, and an
approved file is placed in the submissions directory. **Committing a file into
that directory is the approval step** — nothing is processed until you do.

For each file this source:
  * extracts the document text for the shared extractor, and
  * (on successful extraction) copies the file into the published site under
    ``docs/cfps/`` so it becomes the conference's durable link, and removes the
    original from the submissions directory so it isn't processed twice.

The hosted file's relative link is passed into the extractor as the source URL,
so the resulting row links to the CFP we host (unless the document itself names
a better official homepage, which the extractor prefers).
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
from typing import Iterator, List

from .base import SourceDocument

_TEXT_EXT = {".txt", ".text", ".md"}
_SKIP = {".gitkeep"}


def _slug(stem: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "-", stem).strip("-").lower()
    return s or "cfp"


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(path: str) -> str:
    """Best-effort plain text from a submitted CFP file (pdf/docx/txt/md)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".docx":
            return _extract_docx(path)
        if ext in _TEXT_EXT:
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                return fh.read()
    except Exception as exc:  # a corrupt file shouldn't abort the whole run
        print(f"  ! could not read {os.path.basename(path)}: {exc}")
        return ""
    print(f"  ! unsupported submission type {ext!r} ({os.path.basename(path)})")
    return ""


class SubmissionSource:
    """Treat approved CFP files in a directory as a stream of documents."""

    def __init__(
        self,
        submissions_dir: str = "submissions",
        host_dir: str = "docs/cfps",
        link_prefix: str = "cfps",
    ):
        self.submissions_dir = submissions_dir
        self.host_dir = host_dir
        self.link_prefix = link_prefix

    def _files(self) -> List[str]:
        if not os.path.isdir(self.submissions_dir):
            return []
        out = []
        for name in sorted(os.listdir(self.submissions_dir)):
            if name.startswith(".") or name.lower() in _SKIP:
                continue
            if name.lower() == "readme.md":
                continue
            path = os.path.join(self.submissions_dir, name)
            if os.path.isfile(path):
                out.append(path)
        return out

    def iter_documents(self) -> Iterator[SourceDocument]:
        for path in self._files():
            text = extract_text(path)
            if not text.strip():
                continue  # leave the file in place to inspect/retry
            with open(path, "rb") as fh:
                digest = hashlib.sha1(fh.read()).hexdigest()[:8]
            stem, ext = os.path.splitext(os.path.basename(path))
            hosted_name = f"{_slug(stem)}-{digest}{ext.lower()}"
            link = f"{self.link_prefix}/{hosted_name}"

            marker = {"done": False}

            def on_success(_m=marker):
                _m["done"] = True

            # Hand the extractor the text plus the link we WILL host it at, so
            # the conference row points to the CFP (webpage_source does the same).
            yield SourceDocument(
                text=f"Source URL: {link}\n\nSubmitted call-for-papers document:\n\n{text}",
                origin=f"submission:{os.path.basename(path)}",
                on_success=on_success,
            )

            if marker["done"]:
                # Publish the file and remove it from the approval inbox.
                os.makedirs(self.host_dir, exist_ok=True)
                shutil.copy2(path, os.path.join(self.host_dir, hosted_name))
                os.remove(path)
                print(f"    published {link}")
