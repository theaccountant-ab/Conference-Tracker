"""Process Tally form submission notification emails.

When someone submits the conference form on the website, Tally sends a
notification to the mailbox containing the conference name and a download
link for the uploaded call-for-papers file (PDF or DOCX). This source
detects those emails, downloads the file, extracts its text, and yields a
SourceDocument for the extractor — exactly like any other source.
"""

from __future__ import annotations

import email
import imaplib
import io
import re
import urllib.request
from email.message import Message
from html.parser import HTMLParser
from typing import Iterator, List, Optional, Tuple

from ..config import MailboxConfig
from .base import SourceDocument
from .email_source import _decode, _html_to_text

TALLY_SENDER = "notifications@tally.so"

_USER_AGENT = "ConferenceTracker/0.1 (+https://github.com/theaccountant-ab)"


class _LinkExtractor(HTMLParser):
    """Collect visible text and href links from an HTML email."""

    def __init__(self) -> None:
        super().__init__()
        self._chunks: List[str] = []
        self._links: List[str] = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip = True
        if tag == "a":
            href = dict(attrs).get("href", "")
            if href:
                self._links.append(href)

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            self._skip = False

    def handle_data(self, data):
        if not self._skip and data.strip():
            self._chunks.append(data)

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "\n".join(self._chunks)).strip()

    def links(self) -> List[str]:
        return self._links


def _extract_file_url(html: str) -> Optional[str]:
    """Return the first file-download link found in a Tally notification email."""
    parser = _LinkExtractor()
    try:
        parser.feed(html)
    except Exception:
        return None
    for link in parser.links():
        lower = link.lower().split("?")[0]
        if any(x in lower for x in ("tally.so", ".pdf", ".docx", ".doc")):
            return link
    return None


def _fetch_bytes(url: str) -> bytes:
    req = urllib.request.Request(url, headers={"User-Agent": _USER_AGENT})
    with urllib.request.urlopen(req, timeout=60) as resp:
        return resp.read()


def _pdf_to_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _docx_to_text(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paras)


def _file_to_text(data: bytes, url: str) -> str:
    """Convert downloaded PDF or DOCX bytes to plain text."""
    base = url.lower().split("?")[0]
    if base.endswith(".pdf"):
        return _pdf_to_text(data)
    if base.endswith((".docx", ".doc")):
        return _docx_to_text(data)
    # Unknown extension — try PDF first, then DOCX.
    try:
        text = _pdf_to_text(data)
        if text.strip():
            return text
    except Exception:
        pass
    return _docx_to_text(data)


class TallyEmailSource:
    """IMAP source restricted to Tally form-submission notification emails."""

    def __init__(self, config: MailboxConfig):
        self.config = config

    def _connect(self) -> imaplib.IMAP4:
        cfg = self.config
        conn: imaplib.IMAP4 = (
            imaplib.IMAP4_SSL(cfg.host, cfg.port)
            if cfg.use_ssl
            else imaplib.IMAP4(cfg.host, cfg.port)
        )
        conn.login(cfg.username, cfg.password)
        return conn

    def _disconnect(self, conn: imaplib.IMAP4) -> None:
        for step in (conn.close, conn.logout):
            try:
                step()
            except Exception:
                pass

    def _mark_seen(self, uids: List[bytes]) -> None:
        try:
            conn = self._connect()
        except Exception:
            return
        try:
            conn.select(self.config.folder)
            for uid in uids:
                try:
                    conn.uid("STORE", uid, "+FLAGS", "\\Seen")
                except Exception:
                    pass
        finally:
            self._disconnect(conn)

    def _process(self, msg: Message, uid: bytes) -> Optional[SourceDocument]:
        subject = str(email.header.make_header(
            email.header.decode_header(msg.get("Subject", ""))
        ))
        origin = f"tally:{msg.get('Message-ID', uid.decode())}"

        plain_parts: List[str] = []
        html_raw: Optional[str] = None

        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                disp = str(part.get("Content-Disposition", ""))
                if "attachment" in disp:
                    continue
                payload = part.get_payload(decode=True)
                if payload is None:
                    continue
                text = _decode(payload, part.get_content_charset())
                if ctype == "text/plain":
                    plain_parts.append(text)
                elif ctype == "text/html" and html_raw is None:
                    html_raw = text
        else:
            payload = msg.get_payload(decode=True)
            if payload is not None:
                text = _decode(payload, msg.get_content_charset())
                if msg.get_content_type() == "text/html":
                    html_raw = text
                else:
                    plain_parts.append(text)

        email_text = "\n\n".join(plain_parts)
        if not email_text and html_raw:
            email_text = _html_to_text(html_raw)
        email_text = f"Subject: {subject}\n\n{email_text}".strip()

        file_text = ""
        if html_raw:
            file_url = _extract_file_url(html_raw)
            if file_url:
                try:
                    data = _fetch_bytes(file_url)
                    file_text = _file_to_text(data, file_url)
                    print(f"    Downloaded CFP ({len(data):,} bytes) from {file_url}")
                except Exception as exc:
                    print(f"    ! Could not download CFP from {file_url}: {exc}")

        combined = (
            f"{email_text}\n\n--- Call for Papers ---\n\n{file_text}"
            if file_text.strip()
            else email_text
        )
        return SourceDocument(text=combined, origin=origin) if combined.strip() else None

    def iter_documents(self) -> Iterator[SourceDocument]:
        cfg = self.config
        if not cfg.host or not cfg.username:
            raise RuntimeError(
                "Mailbox is not configured (set CT_MAIL_HOST / CT_MAIL_USER, etc.)"
            )

        conn = self._connect()
        raw_messages: List[Tuple[bytes, bytes]] = []
        try:
            conn.select(cfg.folder)
            criterion = "UNSEEN" if cfg.unseen_only else "ALL"
            typ, data = conn.uid("SEARCH", None, criterion)
            if typ == "OK" and data and data[0]:
                for uid in data[0].split():
                    typ2, msg_data = conn.uid("FETCH", uid, "(RFC822)")
                    if typ2 != "OK" or not msg_data or not msg_data[0]:
                        continue
                    msg = email.message_from_bytes(msg_data[0][1])
                    if TALLY_SENDER in msg.get("From", ""):
                        raw_messages.append((uid, msg_data[0][1]))
        finally:
            self._disconnect(conn)

        processed_uids: List[bytes] = []
        try:
            for uid, raw in raw_messages:
                msg = email.message_from_bytes(raw)
                doc = self._process(msg, uid)
                if doc is None:
                    continue
                marker = {"done": False}

                def on_success(_m=marker):
                    _m["done"] = True

                yield SourceDocument(text=doc.text, origin=doc.origin, on_success=on_success)
                if marker["done"]:
                    processed_uids.append(uid)
        finally:
            if cfg.unseen_only and processed_uids:
                self._mark_seen(processed_uids)
